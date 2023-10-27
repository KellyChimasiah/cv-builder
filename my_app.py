from docx import Document
from docx.shared import Inches
import pyttsx3

# Initialize the text-to-speech engine
engine = pyttsx3.init()

def speak(text):
    engine.say(text)
    engine.runAndWait()

document = Document()

document.add_picture('kelly4.jpg', width=Inches(2.0))
# Name, Phone Number, Email
speak('What is your name?')
name = input('What is your name? ')
speak('Hello ' + name + ', how are you today?')

speak(name + ', what is your phone number?')
phone_number = input('What is your phone number? ')

speak('What is your emai?')
email = input('What is your email? ')

data = f"NAME: {name}\nPHONE NUMBER: {phone_number}\nEMAIL: {email}"
document.add_paragraph(data)

# About me
document.add_heading('About me')
speak('Tell me about yourself ' + name)
About_me = input('Tell me about yourself: ')
document.add_paragraph(About_me)

# Work Experience
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Name of the company you worked for')
start = input('Date you started working')
to = input('Date you stopped working')

p.add_run(company + '\n').bold = True
p.add_run(start)
p.add_run(to + '\n')

experience_details = input('Describe your work experience at ' + company + ': ')
p.add_run(experience_details)

# More experiences
while True:
    company = input('Name of the company you worked for (or type "done" to finish): ')

    if company.lower() == 'done':
        break

    start = input('Date you started working: ')
    to = input('Date you stopped working: ')

    p.add_run(company + '\n').bold = True
    p.add_run(start + ' - ' + to + '\n')

    experience_details = input('Describe your work experience at ' + company + ': ')
    p.add_run(experience_details + '\n')

    # Skills
    document.add_heading('Skills')
    p = document.add_paragraph()

    while True:
        skill = input('Name a skill you have (or type "done" to finish): ')

        if skill.lower() == 'done':
            break

        p.add_run(skill + '\n')

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Made by Kelly Indeche'

# Save the document
document.save('cv.docx')