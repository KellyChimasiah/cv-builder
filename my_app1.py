from docx import Document
from docx.shared import Inches
import pyttsx3

def speak (text):
    pyttsx3.speak(text)

document = Document()

document.add_picture('kelly4.jpg', width=Inches(2.0))
# Name, Phone Number, Email
name = input ('What is your name?')
speak('Hello' + name + 'how are you today?')

speak(name + 'What is your phone number?' )
phone_number = input ('What is your phone number?')
email = input ('What is your email?')

data = f"NAME: {name}\nPHONE NUMBER: {phone_number}\nEMAIL: {email}"
document.add_paragraph(data)

# About me
document.add_heading('About me')
About_me = input ('Tell me about yourself')
document.add_paragraph(About_me)

# Work Experience
document.add_heading('Work experience')
p=document.add_paragraph()

company =input('Name of the company you worked for')
start =input('Date you started working')
to = input('Date you stoped working')

p.add_run(company + '\n').bold= True
p.add_run(start)
p.add_run(to + '\n')

experience_details=input(
    'Describe your work experience at '+ company
)
p.add_run(experience_details)

#more experiences
while True:
    company = input('Name of the company you worked for (or type "done" to finish): ')
    
    if company.lower() == 'done':
        break
    
    start = input('Date you started working: ')
    to = input('Date you stopped working: ')
    
    p.add_run(company + '\n').bold = True
    p.add_run(start + ' - ' + to + '\n')
    
    experience_details = input('Describe your work experience at ' + company + ': ')
    p.add_run(experience_details + '\n')

    #skills
    document.add_heading('Skills')
    p= document.add_paragraph()

    skill=input('Name the Skills you have')
    p.style = 'List Bullet'

    while True:
        skill = input('Skills you have "If done type done"')

        if skill.lower == 'done':
            break
        skill=input('Name the Skills you have')
    p.style = 'List Bullet'

    #footer
    section = document.sections[0]
    footer = section.footer
    p.footer.paragraphs[0]
    p.text = 'Made by Kelly Indeche'

document.save('cv.docx')